from django.shortcuts import render
from django.contrib.auth import authenticate,login

from django.contrib.auth.models import User
from .models import Profile,UploadFile,Bin
from .serializers import RealizerSerializer,LoginSerializer,ProfileSerializer,UploadFileSerializer,BinSerializer

from rest_framework import viewsets
from rest_framework.views import APIView
from rest_framework import serializers
from rest_framework import status
from rest_framework.response import Response
from rest_framework_simplejwt.views import TokenObtainPairView

from rest_framework_simplejwt.tokens import RefreshToken
from rest_framework.permissions import IsAuthenticated
import requests
import PyPDF2
from pathlib import Path
from docx import Document  # For handling .docx files
import csv  # For handling .csv files
import openpyxl  # For handling .xlsx files
from django.http import JsonResponse
import os
import fitz
from fuzzywuzzy import fuzz
import numpy as np
import spacy
from sentence_transformers import SentenceTransformer
from django.shortcuts import get_object_or_404 
import docx  
from sklearn.metrics.pairwise import cosine_similarity

# Create your views here.

class MoveToBinAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, file_id):
        try:
            profile=Profile.objects.get(username_id=request.user.id)
            file = UploadFile.objects.get(id=file_id, postUser=profile)
            
            # Move the file to the bin
            Bin.objects.create(postUserB=profile,fileB=file.file)
            
            # Delete the file from UploadFile
            file.delete()

            return Response({"message": "File moved to bin."}, status=status.HTTP_200_OK)

        except UploadFile.DoesNotExist:
            return Response({"error": "File not found."}, status=status.HTTP_404_NOT_FOUND)
        
class GetBinFilesAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        profile = Profile.objects.get(username_id=request.user.id)
        bin_files = Bin.objects.filter(postUserB=profile)
        serializer = BinSerializer(bin_files, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
        
class SaveFileAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, file_id):
        profile=Profile.objects.get(username_id=request.user.id)
        try:
            file = UploadFile.objects.get(id=file_id, postUser=profile)
            file.saved = True  # Mark the file as saved
            file.save()

            return Response({"message": "File marked as saved."}, status=status.HTTP_200_OK)

        except UploadFile.DoesNotExist:
            return Response({"error": "File not found."}, status=status.HTTP_404_NOT_FOUND)
        
class SavedFilesAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        profile = Profile.objects.get(username_id=request.user.id)  # Get the user's profile
        saved_files = UploadFile.objects.filter(postUser=profile, saved=True)  # Fetch saved files
        serializer = UploadFileSerializer(saved_files, many=True)  # Serialize the data
        return Response(serializer.data, status=status.HTTP_200_OK)

class PostFile(APIView):
    permission_classes=[IsAuthenticated]
    serializer_class=UploadFileSerializer

    def post(self,request):
        profile=Profile.objects.get(username_id=request.user.id)
        
        ##print("post profile:",profile)
        serializer=UploadFileSerializer(data=request.data)
        ##print("serializer:",serializer)
        if serializer.is_valid(raise_exception=True):
                ##print("serializer 1:",serializer)
                upload_file=serializer.save(postUser=profile)
                ##self.process_file(upload_file)
                return Response(serializer.data,status=status.HTTP_200_OK)
        else:
                ##print("serializer 2:",serializer)
                return Response(serializer.errors,status=status.HTTP_400_BAD_REQUEST)
            
model = SentenceTransformer('all-MiniLM-L6-v2')


class GetAllFile(APIView):
    permission_classes = [IsAuthenticated]
    serializer_class = UploadFileSerializer

    def get(self, request):
        # Get user profile
        try:
            profile = Profile.objects.get(username_id=request.user.id)
        except Profile.DoesNotExist:
            return Response({"error": "User profile not found"}, status=404)

        # Get all user-uploaded files
        files = UploadFile.objects.filter(postUser=profile)
        serializer = UploadFileSerializer(files, many=True)
        
        # Extract file paths from the database
        file_paths = [file.file.path for file in files]  # Assuming `file` field stores FileField
        file_texts, file_names = self.load_files(file_paths)

        # Compute similarity
        similar_files = self.compute_all_similarities(file_texts, file_names)

        print("similar_files:",similar_files)
        return Response({
            "files": serializer.data,
            "similar_files": similar_files
        })

    def load_files(self, file_paths):
        """ Extract text from different file formats. """
        file_texts = []
        file_names = []

        for file_path in file_paths:
            if not os.path.exists(file_path):
                print(f"Warning: Skipping missing file - {file_path}")
                continue

            text = self.extract_text(file_path)
            if text:
                file_texts.append(text)
                file_names.append(os.path.basename(file_path))

        return file_texts, file_names

    def extract_text(self, file_path):
        """ Extract text from supported file types. """
        try:
            if file_path.endswith(".pdf"):
                return self.extract_text_from_pdf(file_path)
            elif file_path.endswith(".txt"):
                return self.extract_text_from_txt(file_path)
            elif file_path.endswith(".docx"):
                return self.extract_text_from_docx(file_path)
            elif file_path.endswith(".csv"):
                return self.extract_text_from_csv(file_path)
            elif file_path.endswith(".xlsx"):
                return self.extract_text_from_xlsx(file_path)
            else:
                print(f"Skipping unsupported file type: {file_path}")
                return ""
        except Exception as e:
            print(f"Error processing file {file_path}: {e}")
            return ""

    def extract_text_from_pdf(self, file_path):
        """ Extract text from PDFs. """
        text = ""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()

    def extract_text_from_txt(self, file_path):
        """ Extract text from TXT files. """
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()

    def extract_text_from_docx(self, file_path):
        """ Extract text from DOCX files. """
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs]).strip()

    def extract_text_from_csv(self, file_path):
        """ Extract text from CSV files. """
        with open(file_path, 'r', encoding='utf-8') as file:
            reader = csv.reader(file)
            return '\n'.join([' '.join(row) for row in reader]).strip()

    def extract_text_from_xlsx(self, file_path):
        """ Extract text from XLSX files. """
        wb = openpyxl.load_workbook(file_path)
        sheet = wb.active
        text = ""
        for row in sheet.iter_rows(values_only=True):
            text += ' '.join([str(cell) for cell in row if cell is not None]) + '\n'
        return text.strip()

    def compute_all_similarities(self, file_texts, file_names, threshold=0.8):
        """ Compute pairwise similarities between all files. """
        if not file_texts:
            return []

        file_embeddings = model.encode(file_texts)  # Get embeddings
        similarity_matrix = cosine_similarity(file_embeddings)

        similar_files = []
        for i in range(len(file_names)):
            for j in range(i + 1, len(file_names)):  # Avoid redundant comparisons
                similarity_score = similarity_matrix[i][j]
                if similarity_score >= threshold:
                    similar_files.append({
                        "file1": file_names[i],
                        "file2": file_names[j],
                        "similarity_score": round(similarity_score * 100, 2)
                    })

        return similar_files
   
# 🔹 Helper function to extract text from PDF
def extract_pdf_text(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text if text else "No text found in PDF."

# 🔹 Helper function to extract text from Word files (.docx)
def extract_docx_text(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text if text else "No text found in DOCX."

# 🔹 Helper function to extract data from Excel files (.xlsx)
def extract_xlsx_data(file_path):
    workbook = openpyxl.load_workbook(file_path)
    data = {}
    for sheet in workbook.sheetnames:
        sheet_data = []
        worksheet = workbook[sheet]
        for row in worksheet.iter_rows(values_only=True):
            sheet_data.append(row)
        data[sheet] = sheet_data
    return data if data else "No data found in XLSX." 
  
class GetOneFile(APIView):
     permission_classes=[IsAuthenticated]
     serializer_class=UploadFileSerializer

     def get(self, request, id):
        try:
            # Get the authenticated user's profile
            profile = Profile.objects.get(username_id=request.user.id)

            # Fetch the file object
            file_obj = get_object_or_404(UploadFile, postUser=profile, id=id)

            # Get the file path
            file_path = file_obj.file.path  # Full file system path
            file_extension = os.path.splitext(file_path)[1].lower()

            # Check if the file exists
            if not os.path.exists(file_path):
                return Response({"error": "File not found"}, status=404)

            file_content = None

            # 🔹 Handle different file types
            if file_extension in [".txt", ".csv", ".json", ".log"]:
                # Read plain text files
                with open(file_path, "r", encoding="utf-8") as f:
                    file_content = f.read()

            elif file_extension == ".pdf":
                # Extract text from PDF
                file_content =extract_pdf_text(file_path)

            elif file_extension in [".docx"]:
                # Extract text from Word Document
                file_content =extract_docx_text(file_path)

            elif file_extension in [".xlsx"]:
                # Extract data from Excel file
                file_content =extract_xlsx_data(file_path)

            # Serialize file data and return
            return Response({
                "id": file_obj.id,
                "file_name": file_obj.file.name,
                "file_url": request.build_absolute_uri(file_obj.file.url),
                "file_content": file_content  # May contain text or JSON (for Excel)
            }, status=200)

        except Profile.DoesNotExist:
            return Response({"error": "Profile not found"}, status=404)

        except Exception as e:
            return Response({"error": str(e)}, status=500)

    



import re
class SearchFiles(APIView):
    permission_classes = [IsAuthenticated]
    nlp = spacy.load("en_core_web_md")

    def extract_paragraphs_from_pdf(self,file_path):
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            paragraphs = []
            for page in reader.pages:
                text = page.extract_text()
                # Split text into paragraphs
                paragraphs.extend(text.split("\n\n"))  # Split by double line breaks
            return [para.strip() for para in paragraphs if para.strip()]

    def extract_text_from_txt(self, file_path):
        """Extract text from a plain text file (.txt)"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def extract_text_from_docx(self, file_path):
        """Extract text from a Word document (.docx)"""
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def extract_text_from_csv(self, file_path):
        """Extract text from a CSV file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            reader = csv.reader(file)
            return '\n'.join([' '.join(row) for row in reader])

    def extract_text_from_xlsx(self, file_path):
        """Extract text from an Excel file"""
        wb = openpyxl.load_workbook(file_path)
        sheet = wb.active
        text = ""
        for row in sheet.iter_rows(values_only=True):
            text += ' '.join([str(cell) for cell in row]) + '\n'
        return text
    
    def clean_text(self, text):
        """Normalize the text by removing unwanted characters and extra spaces."""
        # Remove bullet points and special characters (like `•`)
        cleaned_text = re.sub(r"[•\n\r\t]+", " ", text)  # Replace bullet points, newlines, tabs with space
        cleaned_text = re.sub(r"\s{2,}", " ", cleaned_text)  # Replace multiple spaces with single space
        return cleaned_text.strip()
    

    def get(self, request):
        search_text = request.query_params.get('q', '')  
        print("search_text:", search_text)
        print("search_text list:",type(search_text))
        search_text = self.clean_text(search_text)
        print("search_text after:",search_text)
        '''
        if "•" in search_text:
             bullet_index = search_text.find("•")
             search_text=search_text[:bullet_index].strip()
        '''

        profile = Profile.objects.get(username_id=request.user.id)
        # Convert the search text into a spaCy Doc
        #search_doc = self.nlp(search_text.lower())
        #print("search_doc:",search_doc)
        # Get all files associated with the profile
        files = UploadFile.objects.filter(postUser=profile)
        matching_results = []  # Store file and matching sentences
        
        for file_instance in files:
            file_path = Path(file_instance.file.path)  
            print("file_path:", file_path)
            try:
                file_text = ""
                # Check file extension and extract text accordingly
                if file_path.suffix.lower() == '.pdf':
                    '''
                    reader = PdfReader(str(file_path))
                    for page in reader.pages:
                        file_text += page.extract_text()
                    '''
                    paragraphs = self.extract_paragraphs_from_pdf(file_path)
                    file_text = "\n".join(paragraphs)
                elif file_path.suffix.lower() == '.txt':
                    file_text = self.extract_text_from_txt(file_path)
                elif file_path.suffix.lower() == '.docx':
                    file_text = self.extract_text_from_docx(file_path)
                elif file_path.suffix.lower() == '.csv':
                    file_text = self.extract_text_from_csv(file_path)
                elif file_path.suffix.lower() == '.xlsx':
                    file_text = self.extract_text_from_xlsx(file_path)
                else:
                    print(f"Skipping unsupported file type: {file_path.suffix}")
                    continue
                cleaned_file_text = self.clean_text(file_text)
                # Process the extracted text with spaCy and search for sentences
                file_doc = self.nlp(cleaned_file_text)
                #print("file_doc:",file_doc)
                a=re.search(re.escape(search_text.lower()), file_doc.text.lower())
                print("a:",a)
                matching_sentences = [
                    sent.text.strip()
                    for sent in file_doc.sents  # Tokenize text into sentences
                    if search_text.lower() in sent.text.lower()
                ]
                ##Need to improve it and add ctrl+f in react code
                if len(matching_sentences) == 0 and a:
                    matched_text = a.group()  # Extract the matched text
                    matching_sentences.append(matched_text.strip())
                print("matching_sentences:",matching_sentences)
                if matching_sentences:
                    matching_results.append({
                        "file_name": file_instance.file.name,
                        "matching_sentences": matching_sentences
                    })
                print("matching_results:",matching_results)
              
            except Exception as e:
                print(f"Error reading file {file_path}: {e}")

        # Return matching sentences
        return Response({
            "status": "success",
            "message": "Search completed successfully.",
            "data": matching_results
        })



class GetSingleProfile(APIView):
      permission_classes=[IsAuthenticated]
      serializer_class=ProfileSerializer
      
      def get(self,request):
            user=User.objects.get(username=self.request.user)
            profile=Profile.objects.get(username_id=request.user.id)
            print(profile)
            serializer=ProfileSerializer(profile,many=False)
            return Response(serializer.data)
      
class RegisterView(viewsets.ModelViewSet):
    queryset=User.objects.all()
    serializer_class=RealizerSerializer

class LoginUser(APIView):
      
      def post(self,request,format=None):
            serializer =LoginSerializer(data=self.request.data)
            serializer.is_valid(raise_exception=True)
            user = serializer.validated_data['user']
            password = serializer.validated_data['password']
            user=authenticate(username=user,password=password)
            print("is authenticated")
            login(request, user)
            d={}
            
            if user:#Create a field in profile for token. And add auto update logic from backend.
                  t=RefreshToken.for_user(user)
                  #CHeck its working
                  p=Profile.objects.get(username=user.id)
                  print("Profile:",p)
                  p.access_token=str(t.access_token),
                  p.save()
                 
                
                 
                  
                  if p.access_token is not None:
                        print("Access token")
                  else:
                        print("No access token yet")
                        
                  
                  d={"user":user.username,
                  'refresh': str(t),
                  'access': str(t.access_token),
                  }
                  print(d)
                  return Response(d, status=status.HTTP_202_ACCEPTED)
            
            return Response(serializer.errors,status=status.HTTP_400_BAD_REQUEST)
  
      def get(self,request):
            #u=User.objects.all()
            #print(u)
            user = User.objects.get(username=self.request.user)
            serializer=LoginSerializer(user,many=False)
            t=RefreshToken.for_user(user)
            
            d={"user":user.username,
                  'refresh': str(t),
                  'access': str(t.access_token),
                  }
            print(d)
            return Response(d, status=status.HTTP_202_ACCEPTED)
      
class LogoutView(APIView):
      
      permission_classes=(IsAuthenticated,)
      def post(self, request):
            
        try:
            print("Inside log out")
            refresh_token = request.data["refresh_token"]
            print("Refresh:",refresh_token)
            print(request.user)
            #RefreshToken.for_user(user)
            token = RefreshToken.for_user(request.user)
            print("Token:",token)
            token.blacklist()

            return Response(status=status.HTTP_205_RESET_CONTENT)
        except Exception as e:
            print("Inside log out except")
            return Response(status=status.HTTP_400_BAD_REQUEST)

